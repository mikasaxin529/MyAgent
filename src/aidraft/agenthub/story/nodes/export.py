"""export 节点：分镜确认后导出交付四件套。

产物（session 目录下）：
- <片名>_剧本.docx    剧本（场景/动作/台词，docx 用 python-docx 生成）
- <片名>_分镜表.xlsx  分镜表（场/镜/景别/运镜/画面/台词/音效，openpyxl）
- <片名>_预览.html    HTML 预览（分镜卡片流，含立绘图与 image_prompt）
- 立绘图片包          assets/characters/*.png（gen_portraits 已落盘）

跨轮兜底：路由直跳 export 时本轮 state 只有 params，synopsis/characters/
storyboard 都从盘上 state.json 恢复。

实现纯标准库 + python-docx/openpyxl（web 依赖组里已有），不再起子进程——
交付物是结构化文档不是渲染管线，直接在节点里写。
"""
from __future__ import annotations

import html
from typing import Callable

from ..state import (
    StoryState,
    _load_state,
    _session_dir,
    _session_name,
    _step,
)

_MIME = {
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".html": "text/html",
}


def _files_payload(session: str, session_dir):
    """收集 session 目录下的交付文件清单（files 帧用）。"""
    files = []
    for fp in sorted(session_dir.iterdir()):
        if not fp.is_file() or fp.suffix not in _MIME:
            continue
        files.append({
            "name": fp.name,
            "path": f"/files/story/{session}/{fp.name}",
            "size": fp.stat().st_size,
            "mime": _MIME[fp.suffix],
        })
    return files


def _export_docx(path, synopsis: dict, storyboard: dict, characters: dict) -> None:
    """剧本 docx：片头信息 → 角色表 → 按场写动作/台词。"""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    title = synopsis.get("title", "未命名")
    doc.add_heading(f"{title} · 剧本", level=0)
    meta_line = (f"Logline：{synopsis.get('logline', '')}\n"
                 f"主题：{'、'.join(synopsis.get('themes') or [])}")
    p = doc.add_paragraph(meta_line)
    p.runs[0].font.size = Pt(10.5)

    doc.add_heading("角色", level=1)
    for c in characters.get("characters") or []:
        doc.add_paragraph(
            f"{c.get('name', '')}（{c.get('role', '')}）：{c.get('description', '')}",
            style="List Bullet")
        # 戏剧字段行（Sudowrite story bible 式：want/need/arc/voice）
        drama_bits = []
        for label, key in (("目标", "want"), ("课题", "need"),
                           ("弧线", "arc"), ("口吻", "voice"),
                           ("关系", "relationships")):
            if str(c.get(key) or "").strip():
                drama_bits.append(f"{label}：{c[key]}")
        if drama_bits:
            p = doc.add_paragraph("　　" + "；".join(drama_bits))
            for r in p.runs:
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("正文", level=1)
    for sc in storyboard.get("scenes") or []:
        doc.add_heading(
            f"第{sc.get('scene_no', '?')}场　{sc.get('slug', '')}", level=2)
        if sc.get("synopsis"):
            p = doc.add_paragraph(str(sc.get("synopsis")))
            for r in p.runs:
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        for sh in sc.get("shots") or []:
            head = (f"[{sh.get('id', '')}] {sh.get('shot_size', '')}"
                    f"·{sh.get('camera', '')}"
                    + (f"·{sh['camera_angle']}" if sh.get("camera_angle") else "")
                    + (f"·{sh['duration_sec']}s" if sh.get("duration_sec") else "")
                    + (f"→{sh['transition']}" if sh.get("transition") else ""))
            doc.add_paragraph(head).runs[0].bold = True
            action = str(sh.get("action") or sh.get("subject") or "")
            if action:
                doc.add_paragraph(f"　　{action}")
            dlg = str(sh.get("dialogue") or "").strip()
            if dlg:
                doc.add_paragraph(f"　　台词：{dlg}")
            if sh.get("sfx"):
                doc.add_paragraph(f"　　音效：{sh['sfx']}")
    doc.save(str(path))


def _export_xlsx(path, storyboard: dict) -> None:
    """分镜表 xlsx：一行一镜。"""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    wb = Workbook()
    ws = wb.active
    ws.title = "分镜表"
    headers = ["场号", "场景", "镜号", "景别", "运镜", "机位角度", "时长(s)",
               "转场", "画面主体", "动作", "台词", "音效",
               "画面描述(image_prompt)"]
    ws.append(headers)
    head_fill = PatternFill("solid", fgColor="7C5CBF")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = head_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for sc in storyboard.get("scenes") or []:
        for sh in sc.get("shots") or []:
            ws.append([
                sc.get("scene_no", ""), sc.get("slug", ""),
                sh.get("id", ""), sh.get("shot_size", ""),
                sh.get("camera", ""), sh.get("camera_angle", ""),
                sh.get("duration_sec", ""), sh.get("transition", ""),
                sh.get("subject", ""),
                sh.get("action", ""), sh.get("dialogue", ""),
                sh.get("sfx", ""), sh.get("image_prompt", ""),
            ])
    widths = [6, 22, 8, 8, 8, 9, 8, 8, 28, 30, 30, 16, 44]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + i)].width = w
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    wb.save(str(path))


def _export_html(path, synopsis: dict, storyboard: dict,
                 characters: dict, session: str) -> None:
    """HTML 预览：分镜卡片流（立绘图 + 画面描述 + 台词）。"""

    def esc(s) -> str:
        return html.escape(str(s or ""))

    cards = []
    for sc in storyboard.get("scenes") or []:
        shots = []
        for sh in sc.get("shots") or []:
            dlg = str(sh.get("dialogue") or "").strip()
            meta_bits = [str(x) for x in
                         (sh.get("shot_size"), sh.get("camera"),
                          sh.get("camera_angle")) if x]
            if sh.get("duration_sec"):
                meta_bits.append(f"{sh['duration_sec']}s")
            if sh.get("transition"):
                meta_bits.append(f"→{sh['transition']}")
            shots.append(f"""
        <div class="shot">
          <div class="shot-head">
            <span class="shot-id">{esc(sh.get('id'))}</span>
            <span class="shot-meta">{esc(' · '.join(meta_bits))}</span>
          </div>
          <div class="img-prompt">{esc(sh.get('image_prompt'))}</div>
          <div class="action">{esc(sh.get('action') or sh.get('subject'))}</div>
          {f'<div class="dialogue">“{esc(dlg)}”</div>' if dlg else ''}
          {f'<div class="sfx">♪ {esc(sh.get("sfx"))}</div>' if sh.get("sfx") else ''}
        </div>""")
        cards.append(f"""
      <section class="scene">
        <header><b>第{esc(sc.get('scene_no'))}场</b> {esc(sc.get('slug'))}
          <span class="scene-syn">{esc(sc.get('synopsis'))}</span></header>
        {''.join(shots)}
      </section>""")

    chars_html = []
    for c in characters.get("characters") or []:
        portrait = ""
        if c.get("portrait"):
            portrait = (f'<img class="portrait" src="{esc(c["portrait"])}" '
                        f'alt="{esc(c.get("name"))}" loading="lazy">')
        drama = ""
        if any(str(c.get(k) or "").strip() for k in
               ("want", "need", "arc", "voice", "relationships")):
            drama = (f'<p class="drama">🎯 {esc(c.get("want"))}　'
                     f'🧩 {esc(c.get("need"))}<br>'
                     f'📈 {esc(c.get("arc"))}　🗣 {esc(c.get("voice"))}<br>'
                     f'🔗 {esc(c.get("relationships"))}</p>')
        chars_html.append(f"""
        <div class="char">{portrait}
          <div><b>{esc(c.get('name'))}</b>
            <span class="role">{esc(c.get('role'))}</span>
            <p>{esc(c.get('description'))}</p>
            {drama}
          </div>
        </div>""")

    page = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{esc(synopsis.get('title'))} · 分镜预览</title>
<style>
:root {{ --ink:#2B2440; --sub:#8A82A6; --accent:#7C5CBF; --bg:#FAF9FC; --card:#fff; }}
* {{ box-sizing:border-box; }}
body {{ margin:0; background:var(--bg); color:var(--ink);
  font:14px/1.7 "PingFang SC","Microsoft YaHei",sans-serif; }}
.wrap {{ max-width:960px; margin:0 auto; padding:32px 20px 64px; }}
h1 {{ font-size:26px; margin:0 0 4px; }}
.logline {{ color:var(--sub); margin:0 0 24px; }}
.scene {{ background:var(--card); border-radius:12px; padding:16px 18px;
  margin:16px 0; box-shadow:0 1px 4px rgba(60,40,120,.08); }}
.scene header {{ border-bottom:1px solid #EEE9F5; padding-bottom:8px; margin-bottom:10px; }}
.scene-syn {{ display:block; color:var(--sub); font-size:12.5px; margin-top:2px; }}
.shot {{ padding:10px 0; border-bottom:1px dashed #F0ECF7; }}
.shot:last-child {{ border-bottom:0; }}
.shot-head {{ display:flex; gap:10px; align-items:center; }}
.shot-id {{ font:600 12px ui-monospace,monospace; color:var(--accent);
  background:#F1EBFA; border-radius:4px; padding:1px 7px; }}
.shot-meta {{ color:var(--sub); font-size:12px; }}
.img-prompt {{ color:#5C5478; font-size:12.5px; margin-top:4px; }}
.action {{ margin-top:4px; }}
.dialogue {{ margin-top:4px; font-weight:600; }}
.sfx {{ color:var(--sub); font-size:12px; }}
.chars {{ display:flex; flex-wrap:wrap; gap:14px; margin-bottom:8px; }}
.char {{ display:flex; gap:10px; background:var(--card); border-radius:12px;
  padding:12px; flex:1 1 240px; box-shadow:0 1px 4px rgba(60,40,120,.08); }}
.portrait {{ width:72px; height:72px; object-fit:cover; border-radius:8px;
  background:#F1EBFA; flex:none; }}
.role {{ color:var(--accent); font-size:12px; margin-left:6px; }}
.char p {{ margin:4px 0 0; font-size:12.5px; color:#5C5478; }}
.drama {{ color:#6B5A93; border-top:1px dashed #E5DEF0; padding-top:6px; }}
h2 {{ font-size:16px; margin:28px 0 10px; }}
</style></head><body><div class="wrap">
<h1>{esc(synopsis.get('title'))}</h1>
<p class="logline">{esc(synopsis.get('logline'))}
　主题：{esc('、'.join(synopsis.get('themes') or []))}</p>
<h2>角色</h2><div class="chars">{''.join(chars_html)}</div>
<h2>分镜</h2>
{''.join(cards)}
</div></body></html>"""
    path.write_text(page, encoding="utf-8")


def _make_export_node(emitter: Callable[[dict], None] | None):
    """export 节点工厂：导出 docx/xlsx/html 三件 + 图片包汇总。"""

    async def export(state: StoryState) -> dict:
        _step(emitter, "export", "导出交付", "running")

        visited = list(state.get("nodes_visited") or [])
        if "export" not in visited:
            visited.append("export")

        params = state.get("story_params", {})
        # 跨轮（路由直跳 export）时 synopsis/characters/storyboard 都在盘上，
        # 不在本轮 state 里——统一查盘兜底（与 gen_storyboard 同思路）
        disk = _load_state(params) if params.get("title") else {}
        synopsis = state.get("story_synopsis") or disk.get("story_synopsis") or {}
        characters = state.get("story_characters") or disk.get("story_characters") or {}
        storyboard = state.get("story_storyboard") or disk.get("story_storyboard") or {}
        if not storyboard.get("scenes"):
            _step(emitter, "export", "导出交付", "error", "分镜缺失")
            return {"story_error": "分镜缺失，无法导出",
                    "nodes_visited": visited}

        session = _session_name(params)
        session_dir = _session_dir(params)
        session_dir.mkdir(parents=True, exist_ok=True)
        safe_title = "".join(c for c in session if c.strip()) or "story"

        errors: list[str] = []
        try:
            _export_docx(session_dir / f"{safe_title}_剧本.docx",
                         synopsis, storyboard, characters)
        except Exception as exc:  # noqa: BLE001 - 单件失败不挡其余
            errors.append(f"剧本 docx：{exc}")
        try:
            _export_xlsx(session_dir / f"{safe_title}_分镜表.xlsx", storyboard)
        except Exception as exc:  # noqa: BLE001
            errors.append(f"分镜表 xlsx：{exc}")
        try:
            _export_html(session_dir / f"{safe_title}_预览.html",
                         synopsis, storyboard, characters, session)
        except Exception as exc:  # noqa: BLE001
            errors.append(f"预览 html：{exc}")

        files = _files_payload(session, session_dir)
        if errors and not files:
            err = "；".join(errors)
            _step(emitter, "export", "导出交付", "error", err[:120])
            return {"story_error": err, "nodes_visited": visited}

        detail = f"{len(files)} 个文件已导出"
        if errors:
            detail += f"（部分失败：{'；'.join(e[:40] for e in errors)}）"
        _step(emitter, "export", "导出交付", "done", detail)
        return {"story_files": files,
                "story_error": "；".join(errors),
                "nodes_visited": visited}

    return export
