"""JSON → 教案 .docx 渲染器（python-docx）。

教案 10 模块：
课题 / 教材版本 / 教学目标(competency+dimension双标签) / 教学重难点 /
教学准备 / 课时安排 / 教学过程(分课时,教师活动+学生活动+设计意图) /
板书设计 / 作业设计(分层三级) / 教学反思

配色与字号一律走 design_tokens 主题代理（曾硬编码 RGBColor 与 pptx 漂移，
现统一以 ACTIVE_THEME.pal 为唯一真相）。字体名不主题化（docx 字体与
配色无关，主题化属过度设计）。

退出码：0 成功 / 1 异常 / 2 前置缺失
"""
from __future__ import annotations
import sys
import json
from pathlib import Path

# Windows 控制台默认 GBK，输出 ✓/✗ 等 Unicode 会崩；强制 UTF-8
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass

from common import design_tokens as T
from common.schema import stage_short, validate as _schema_validate

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _rgb(hexstr: str) -> RGBColor:
    return RGBColor.from_string(hexstr)


def _c(key: str, fallback: str = "000000") -> RGBColor:
    """当前主题色 → RGBColor。"""
    return _rgb(T.ACTIVE_THEME.pal.get(key, fallback))


# 模块级常量仅作默认值占位；真正取值发生在 render() 后按当前主题解析。
# 注意：_set_font 等的 color 默认参数不能写死（import 时求值会冻结主题），
# 一律 color=None → 函数内回退到当前主题 TITLE_TEXT。


def _set_font(run, name="微软雅黑", size=12, color=None, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color or _c("TITLE_TEXT")
    run.font.bold = bold
    # 中文字体需单独设 eastAsia
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 教案文档字号档：主题 font_scale 的 d_h1/d_h2/d_h3（stage 由调用方设好）
    role = {1: "d_h1", 2: "d_h2", 3: "d_h3"}.get(level, "d_body")
    run = p.add_run(text)
    _set_font(run, size=T.font_for(_STAGE, role), color=_c("ACCENT"), bold=True)
    p.space_after = Pt(6)
    return p


def _para(doc, text, size=None, color=None, bold=False, indent=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    _set_font(run, size=size if size is not None else T.font_for(_STAGE, "d_body"),
              color=color, bold=bold)
    return p


def _bullet(doc, text, size=None, level=0):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.74 + level * 0.74)
    run = p.add_run(("• " if level == 0 else "○ ") + text)
    _set_font(run, size=size if size is not None else T.font_for(_STAGE, "d_body"),
              color=_c("TITLE_TEXT"))
    return p


def _add_horiz_rule(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), T.ACTIVE_THEME.pal.get("DOCX_RULE", "E8D5C0"))
    pbdr.append(bottom)
    pPr.append(pbdr)


# 当前页学段简称（render() 开头按 doc 设置；字体大小档位查表用）
_STAGE = "中"


def _append_images(doc, document):
    """教案末尾附课件插图（最小实现：src 有效的 image 元素逐个居中插入）。"""
    seen = set()
    added = False
    for sj in doc.get("slides", []):
        for el in sj.get("elements", []):
            if el.get("type") != "image":
                continue
            src = el.get("src", "")
            if not src or src in seen or not Path(src).is_file():
                continue
            seen.add(src)
            if not added:
                _heading(document, "附录：课件配图", 1)
                added = True
            try:
                pic_p = document.add_paragraph()
                pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pic_p.add_run()
                run.add_picture(src, width=Inches(5.5))
                cap = el.get("caption")
                if cap:
                    _para(document, cap,
                          color=_c("TEXT_LIGHT"),
                          size=T.font_for(_STAGE, "d_small"))
                    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                continue  # 插图失败不影响教案正文


def render(doc: dict, out_path: str) -> str:
    """渲染 doc → 教案 .docx。返回文件路径。
    优先用 lessonPlan；缺则从 meta + slides 派生。"""
    global _STAGE
    meta = doc["meta"]
    _STAGE = stage_short(meta.get("stage", "中段"))
    accent = _c("ACCENT")
    text_c = _c("TITLE_TEXT")
    light_c = _c("TEXT_LIGHT")
    lp = doc.get("lessonPlan") or {}
    handout = doc.get("handout") or {}

    document = Document()
    # 页边距
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    base = lp.get("base", {}) if lp else {}
    title = lp.get("title", meta["title"])

    # ---- 课题 ----
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_font(run, size=T.font_for(_STAGE, "d_title"), color=accent, bold=True)

    _add_horiz_rule(document)

    # ---- 基本信息表 ----
    table = document.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    info = [
        ("教材版本", base.get("textbook") or meta.get("textbook", "")),
        ("年级", base.get("grade") or f"{meta['grade']}年级"),
        ("课型", base.get("lessonType") or meta.get("lessonType", "")),
        ("课时", base.get("periods") or f"{meta.get('periods',1)}课时"),
    ]
    for i, (k, v) in enumerate(info):
        r, c = divmod(i, 2)
        c0 = table.cell(r, c*2); c1 = table.cell(r, c*2+1)
        for cell, txt, bold in [(c0, k, True), (c1, v, False)]:
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            _set_font(run, size=T.font_for(_STAGE, "d_small"),
                      color=text_c if not bold else accent, bold=bold)

    document.add_paragraph()

    # ---- 教学目标 ----
    _heading(document, "一、教学目标", 1)
    objectives = lp.get("objectives") or meta.get("objectives", [])
    for obj in objectives:
        if isinstance(obj, dict):
            tag = f"【{obj.get('competency','')}"
            if obj.get("dimension"):
                tag += f"/{obj['dimension']}"
            tag += "】"
            _bullet(document, f"{tag}{obj.get('content','')}")
        else:
            _bullet(document, str(obj))

    # ---- 教学重难点 ----
    _heading(document, "二、教学重难点", 1)
    _para(document, "教学重点：", size=T.font_for(_STAGE, "d_label"), color=accent, bold=True)
    for k in (lp.get("keyPoints") or meta.get("keyPoints", [])):
        _bullet(document, k)
    _para(document, "教学难点：", size=T.font_for(_STAGE, "d_label"), color=accent, bold=True)
    for d in (lp.get("difficulties") or meta.get("difficulties", [])):
        _bullet(document, d)

    # ---- 教学准备 ----
    _heading(document, "三、教学准备", 1)
    _para(document, lp.get("preparation") or "多媒体课件、生字卡片")

    # ---- 课时安排 ----
    _heading(document, "四、课时安排", 1)
    _para(document, f"本课共 {meta.get('periods',1)} 课时。")

    # ---- 教学过程 ----
    _heading(document, "五、教学过程", 1)
    process = lp.get("teachingProcess", [])
    if not process:
        # 缺省：从 slides 派生简版
        for per in sorted(set(s.get("period",1) for s in doc["slides"])):
            _heading(document, f"第 {per} 课时", 2)
            for s in [x for x in doc["slides"] if x.get("period",1)==per]:
                _bullet(document, s.get("title",""))
    else:
        per_groups = {}
        for phase in process:
            ptag = phase.get("period", 1)
            per_groups.setdefault(ptag, []).append(phase)
        if not per_groups:
            per_groups[1] = process
        for per in sorted(per_groups):
            _heading(document, f"第 {per} 课时", 2)
            for phase in per_groups[per]:
                _para(document, phase.get("phase", ""),
                      size=T.font_for(_STAGE, "d_label"), color=accent, bold=True)
                _para(document, f"（{phase.get('duration','')}）",
                      size=T.font_for(_STAGE, "d_small"), color=light_c)
                for act in phase.get("activities", []):
                    _bullet(document, f"教师：{act.get('teacher','')}")
                    _bullet(document, f"学生：{act.get('student','')}", level=1)
                if phase.get("design"):
                    _para(document, phase["design"],
                          size=T.font_for(_STAGE, "d_small"), color=light_c, indent=0.5)

    # ---- 板书设计 ----
    _heading(document, "六、板书设计", 1)
    bd = lp.get("boardDesign", {})
    if isinstance(bd, dict) and bd.get("structure"):
        _para(document, bd["structure"], size=T.font_for(_STAGE, "d_label"))
    else:
        _para(document, "（见课件板书设计页）", color=light_c)

    # ---- 作业设计 ----
    _heading(document, "七、作业设计（分层）", 1)
    hw = lp.get("homework") or handout
    levels = hw.get("levels", []) if isinstance(hw, dict) else []
    if not levels:
        levels = [{"level": "基础", "items": ["抄写生字", "朗读课文"]}]
    tone = T.ACTIVE_THEME.pal.get("TONE", {})
    level_colors = {
        "基础": accent,
        "提升": _rgb(tone.get(3, "5BA88A")),   # 三声绿
        "拓展": _rgb(tone.get(4, "5B8AB5")),   # 四声蓝
    }
    for lv in levels:
        color = level_colors.get(lv.get("level",""), text_c)
        _para(document, f"【{lv.get('level','')}】",
              size=T.font_for(_STAGE, "d_label"), color=color, bold=True)
        for item in lv.get("items", []):
            _bullet(document, item)

    # ---- 教学反思 ----
    _heading(document, "八、教学反思", 1)
    _para(document, lp.get("reflection") or "（课后填写）", color=light_c)

    # ---- 课件配图（可选，src 已由 render_all 解析为绝对路径）----
    _append_images(doc, document)

    document.save(out_path)
    return str(out_path)


def main(argv=None):
    import argparse
    p = argparse.ArgumentParser(description="JSON → 教案 DOCX 渲染")
    p.add_argument("input", help="课程 JSON 文件")
    p.add_argument("-o", "--output", help="输出 .docx 路径")
    args = p.parse_args(argv)
    try:
        with open(args.input, encoding="utf-8") as f:
            doc = json.load(f)
        _schema_validate(doc)
        T.set_theme(doc.get("meta", {}).get("theme"))
    except Exception as e:
        print(f"[docx] ✗ 输入无效：{e}", file=sys.stderr)
        return 2
    out = args.output or str(Path(args.input).with_suffix(".docx"))
    try:
        fp = render(doc, out)
        print(f"[docx] ✓ 已生成：{fp}")
        return 0
    except Exception as e:
        print(f"[docx] ✗ 渲染失败：{e}", file=sys.stderr)
        import traceback; traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
